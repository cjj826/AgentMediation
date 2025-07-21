import http.client
import json
from tqdm import tqdm

def retrieve_articles(case):
    conn = http.client.HTTPConnection("web.megatechai.com", 33615)
    payload = json.dumps({
    "query": case,
    "need_answer": 0
    })
    headers = {
    'Content-Type': 'application/json'
    }
    conn.request("POST", "/case_app/wenshu_search/search_and_answer", payload, headers)
    res = conn.getresponse()
    data = res.read()
    data = json.loads(data.decode("utf-8"))
    wenshu_results = data['data']['wenshu_results']

    articles = []

    for wenshu in wenshu_results:
        xiangguanfatiaos = wenshu['xiangguanfatiao']
        if isinstance(xiangguanfatiaos, list):
            for xiangguanfatiao in xiangguanfatiaos:
                    articles.append(xiangguanfatiao['law_name'] + xiangguanfatiao['xuhao'])
        
    return articles

# datas = open('./case_back/case_back_0506_article_point_deepseek_v3.json', 'r').readlines()
# retrieve_datas = open('./case_back/case_back_0506_article_point_deepseek_v3_retrieve.json', 'a')

# for data in tqdm(datas):
#     data = json.loads(data)
#     case_back = data['basic_back']
#     articles = retrieve_articles(case_back)
#     data['retrieve_articles'] = articles
#     retrieve_datas.write(json.dumps(data, ensure_ascii=False) + '\n')
#     retrieve_datas.flush()

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import json

def json_to_mediation_word(data, output_path) -> None:
    doc = Document()
    
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        run = heading.runs[0]
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '宋体'
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_heading("一、案情背景", level=1)
    background = data.get("global_prompt", {})
    add_paragraph(background)

    add_heading("二、自我介绍", level=1)
    history = data.get("statement_history", [])
    if isinstance(history, list):
        for item in history:
            if isinstance(item, dict):
                speaker = item.get("agent_name") or item.get("speaker") or item.get("name", "未知发言人")
                content = item.get("content") or item.get("statement", "")
                add_paragraph(f"{speaker}：", bold=True)
                add_paragraph(content)

    add_heading("三、初步调解方案", level=1)
    option = data.get("mediation_option", {})
    if option:
        # add_paragraph(f"争议焦点：{option.get('争议焦点', '')}")
        # add_paragraph("解纷依据：", bold=True)
        # for law in option.get("解纷依据", []):
        #     add_paragraph(f"- {law}")
        add_paragraph("调解方案：", bold=True)
        add_paragraph(option.get("调解方案", ""))

    add_heading("四、协商过程", level=1)
    round = 1
    for msg in data.get("bargain_messsages", []):
        speaker = msg.get("agent_name", "未知发言人")
        if speaker.find("调解员") != -1:
            add_heading(f"第{round}轮调解", level=2)
            round += 1
        content = msg.get("content", "")
        add_paragraph(f"{speaker}：", bold=True)
        add_paragraph(content)

    add_heading("五、最终调解方案", level=1)
    final = data.get("final_mediation", {})
    if final:
        add_paragraph("调解方案：", bold=True)
        add_paragraph(final.get("调解方案", ""))

    add_heading("六、调解结果与满意度", level=1)
    for result_block in data.get("result", []):
        acceptance = result_block.get("是否接受", {})
        satisfaction = result_block.get("满意度", {})
        for person, info in acceptance.items():
            add_paragraph(f"{person}是否接受：{info.get('是否接受', '')}")
            add_paragraph(f"理由：{info.get('理由', '')}")
        for person, info in satisfaction.items():
            add_paragraph(f"{person}满意度评级：{info.get('满意度评级', '')}")
            add_paragraph(f"理由：{info.get('理由', '')}")

    # 保存文档
    doc.save(output_path)
    print(f"✅ Word文档已保存至：{output_path}")
